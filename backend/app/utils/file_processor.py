import os
import uuid
import logging
import re
from pathlib import Path
from typing import Dict, Any, List, Tuple

from fastapi import HTTPException, status

logger = logging.getLogger(__name__)


class SecurityScanner:
    """
    Basic security scanner for uploaded files
    """

    # Potentially dangerous file extensions
    DANGEROUS_EXTENSIONS = {'.exe', '.bat', '.cmd', '.com', '.pif', '.scr', '.vbs', '.js', '.jar', '.sh'}

    # Suspicious content patterns
    SUSPICIOUS_PATTERNS = [
        r'password\s*[:=]\s*\w+',
        r'secret\s*[:=]\s*\w+',
        r'api[_-]?key\s*[:=]\s*\w+',
        r'token\s*[:=]\s*\w+',
        r'credit[_-]?card',
        r'ssn\s*\d{3}[-\s]?\d{2}[-\s]?\d{4}',
        r'^\s*<\s*script\s*>',
        r'javascript\s*:',
        r'eval\s*\(',
        r'system\s*\(',
        r'exec\s*\(',
        r'import\s+os',
        r'subprocess\s*\.',
    ]

    MALICIOUS_KEYWORDS = [
        'virus', 'trojan', 'malware', 'ransomware', 'spyware',
        'hack', 'exploit', 'backdoor', 'rootkit',
        'phish', 'scam', 'fraud',
        'bomb', 'attack', 'weapon'
    ]

    @classmethod
    def scan_file_content(cls, content: str, filename: str) -> Dict[str, Any]:
        """
        Scan file content for security issues
        """
        issues = []
        risk_score = 0

        # Check filename for suspicious extensions
        file_extension = Path(filename).suffix.lower()
        if file_extension in cls.DANGEROUS_EXTENSIONS:
            issues.append(f"Suspicious file extension: {file_extension}")
            risk_score += 50

        # Check for suspicious patterns
        for pattern in cls.SUSPICIOUS_PATTERNS:
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                issues.append(f"Potentially sensitive content detected: {pattern}")
                risk_score += 20

        # Check for malicious keywords
        content_lower = content.lower()
        for keyword in cls.MALICIOUS_KEYWORDS:
            if keyword in content_lower:
                issues.append(f"Potentially malicious keyword detected: {keyword}")
                risk_score += 15

        # Check for excessive special characters (might indicate encoding attack)
        special_chars = sum(1 for c in content if ord(c) > 127)
        if special_chars > len(content) * 0.1:
            issues.append("High number of special characters (possible encoding attack)")
            risk_score += 10

        # Determine risk level
        if risk_score >= 70:
            risk_level = "high"
        elif risk_score >= 40:
            risk_level = "medium"
        elif risk_score >= 20:
            risk_level = "low"
        else:
            risk_level = "safe"

        return {
            "is_safe": risk_score < 30,  # Allow low risk content
            "risk_score": risk_score,
            "risk_level": risk_level,
            "issues": issues
        }


class FileProcessor:

    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
    MAX_FILE_SIZE = 50 * 1024 * 1024
    MAX_CHUNK_SIZE = 4000  # Maximum characters per chunk for AI processing
    MIN_CHUNK_SIZE = 500   # Minimum characters per chunk
    @classmethod
    def validate_file(cls, filename: str, file_size: int) -> str:
        if not filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No file provided"
            )

        file_extension = Path(filename).suffix.lower()

        if file_extension not in cls.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type not supported. Allowed types: {', '.join(cls.ALLOWED_EXTENSIONS)}"
            )

        if file_size > cls.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail="File too large. Maximum size: 50MB"
            )

        return file_extension

    @classmethod
    def generate_unique_filename(cls, original_filename: str) -> Tuple[str, Path]:
        file_id = str(uuid.uuid4())
        sanitized_filename = f"{file_id}_{original_filename}"
        upload_dir = Path("uploads/documents")
        upload_dir.mkdir(parents=True, exist_ok=True)

        return sanitized_filename, upload_dir / sanitized_filename

    @classmethod
    def save_file(cls, file_path: Path, content: bytes) -> None:
        try:
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error saving file: {str(e)}"
            )

    @classmethod
    def extract_text_content(cls, file_path: Path, file_extension: str) -> str:
        """
        Extract text content from file with robust error handling
        """
        try:
            if file_extension == ".pdf":
                return cls._extract_pdf_text(file_path)
            elif file_extension in {".docx", ".doc"}:
                return cls._extract_docx_text(file_path)
            elif file_extension == ".txt":
                return cls._extract_txt_text(file_path)
            else:
                return "Unsupported file type for text extraction."
        except ImportError as e:
            return f"Missing required library for {file_extension} processing: {str(e)}"
        except FileNotFoundError as e:
            return f"File not found during processing: {str(e)}"
        except PermissionError as e:
            return f"Permission denied accessing file: {str(e)}"
        except MemoryError as e:
            return f"File too large to process (memory error): {str(e)}"
        except UnicodeDecodeError as e:
            return f"Encoding error in text file: {str(e)}"
        except Exception as e:
            return f"Unexpected error during text extraction: {str(e)}"

    @classmethod
    def _extract_pdf_text(cls, file_path: Path) -> str:
        """
        Extract text from PDF with robust error handling
        """
        try:
            import pdfplumber
            content = ""
            page_count = 0

            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"Processing PDF with {total_pages} pages")

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            content += text.strip() + "\n"
                            page_count += 1
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                        continue

            if not content.strip():
                return "No text content could be extracted from this PDF. The file may be image-based or corrupted."

            logger.info(f"Successfully extracted text from {page_count}/{total_pages} pages")
            return content.strip()

        except ImportError:
            raise ImportError("pdfplumber library not installed")
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")

    @classmethod
    def _extract_docx_text(cls, file_path: Path) -> str:
        """
        Extract text from DOCX with robust error handling
        """
        try:
            from docx import Document
            content = ""
            paragraph_count = 0

            doc = Document(file_path)

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    content += paragraph.text.strip() + "\n"
                    paragraph_count += 1

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            content += cell.text.strip() + " "

            content = content.strip()

            if not content:
                return "No text content could be extracted from this document."

            logger.info(f"Successfully extracted text from DOCX: {paragraph_count} paragraphs")
            return content

        except ImportError:
            raise ImportError("python-docx library not installed")
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")

    @classmethod
    def _extract_txt_text(cls, file_path: Path) -> str:
        """
        Extract text from TXT with robust encoding handling
        """
        encodings_to_try = ["utf-8", "utf-16", "latin-1", "cp1252"]

        for encoding in encodings_to_try:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read().strip()
                    if content:
                        logger.info(f"Successfully read TXT file with {encoding} encoding")
                        return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"Error reading TXT with {encoding}: {str(e)}")
                continue

        return "Could not decode text file content with any supported encoding."

    @classmethod
    def analyze_content(cls, content: str) -> Dict[str, Any]:
        """
        Analyze content quality and characteristics with validation
        """
        if not content or not content.strip():
            return {
                "word_count": 0,
                "difficulty_level": "easy",
                "content_quality": "empty",
                "validation_errors": ["No content extracted"],
                "is_valid": False
            }

        # Basic content validation
        validation_result = cls._validate_content(content)

        if not validation_result["is_valid"]:
            return {
                "word_count": len(content.split()),
                "difficulty_level": "easy",
                "content_quality": "invalid",
                "validation_errors": validation_result["errors"],
                "is_valid": False,
                **validation_result
            }

        # Calculate word count
        word_count = len(content.split())

        # Determine difficulty level
        difficulty_level = "easy" if word_count < 500 else "medium" if word_count < 2000 else "hard"

        # Calculate content quality score
        quality_score = cls._calculate_content_quality(content, word_count)

        return {
            "word_count": word_count,
            "difficulty_level": difficulty_level,
            "content_quality": quality_score["level"],
            "quality_score": quality_score["score"],
            "validation_errors": [],
            "is_valid": True,
            "language_detected": cls._detect_language(content),
            "encoding_issues": validation_result.get("encoding_issues", 0)
        }

    @classmethod
    def _validate_content(cls, content: str) -> Dict[str, Any]:
        """
        Validate extracted content for basic quality checks
        """
        errors = []
        encoding_issues = 0

        # Check for minimum length
        if len(content) < 50:
            errors.append("Content too short (minimum 50 characters required)")

        # Check for excessive special characters (might indicate poor extraction)
        special_chars = sum(1 for c in content if not c.isalnum() and not c.isspace() and c not in '.,!?;:"()-')
        if special_chars > len(content) * 0.1:  # More than 10% special characters
            errors.append("High number of special characters detected (possible extraction issue)")

        # Check for repetitive content (might indicate OCR errors)
        words = content.split()
        if len(words) > 10:
            unique_words = set(words)
            repetition_ratio = len(unique_words) / len(words)
            if repetition_ratio < 0.3:  # Less than 30% unique words
                errors.append("High word repetition detected (possible OCR or extraction issue)")

        # Check for encoding issues (weird characters)
        weird_chars = sum(1 for c in content if ord(c) > 65535 or (ord(c) < 32 and c not in '\n\r\t'))
        if weird_chars > len(content) * 0.01:  # More than 1% weird characters
            encoding_issues = weird_chars

        # Check for minimum word count
        if len(words) < 10:
            errors.append("Insufficient word count (minimum 10 words required)")

        return {
            "is_valid": len(errors) == 0,
            "errors": errors,
            "encoding_issues": encoding_issues
        }

    @classmethod
    def _calculate_content_quality(cls, content: str, word_count: int) -> Dict[str, Any]:
        """
        Calculate content quality score
        """
        score = 0
        max_score = 100

        # Length score (0-30 points)
        if word_count > 1000:
            score += 30
        elif word_count > 500:
            score += 25
        elif word_count > 200:
            score += 20
        elif word_count > 50:
            score += 10
        else:
            score += 0

        # Character diversity score (0-20 points)
        unique_chars = len(set(content.lower()))
        if unique_chars > 50:
            score += 20
        elif unique_chars > 30:
            score += 15
        elif unique_chars > 20:
            score += 10
        else:
            score += 5

        # Sentence structure score (0-30 points)
        sentences = [s.strip() for s in content.split('.') if s.strip()]
        avg_sentence_length = sum(len(s.split()) for s in sentences) / len(sentences) if sentences else 0

        if 10 <= avg_sentence_length <= 25:
            score += 30
        elif 5 <= avg_sentence_length <= 30:
            score += 20
        else:
            score += 10

        # Coherence score (0-20 points) - basic check for paragraph structure
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        if len(paragraphs) > 3:
            score += 20
        elif len(paragraphs) > 1:
            score += 15
        else:
            score += 5

        # Determine quality level
        if score >= 80:
            level = "excellent"
        elif score >= 60:
            level = "good"
        elif score >= 40:
            level = "fair"
        else:
            level = "poor"

        return {
            "score": score,
            "level": level,
            "max_score": max_score
        }

    @classmethod
    def _detect_language(cls, content: str) -> str:
        """
        Basic language detection (very simple heuristic)
        """
        # This is a very basic implementation
        # In production, you'd use a proper library like langdetect
        english_words = sum(1 for word in content.lower().split() if word in {
            'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'
        })

        total_words = len(content.split())
        english_ratio = english_words / total_words if total_words > 0 else 0

        if english_ratio > 0.1:
            return "english"
        else:
            return "unknown"

    @classmethod
    def chunk_text(cls, text: str, max_chunk_size: int = None, min_chunk_size: int = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks suitable for AI processing
        """
        if max_chunk_size is None:
            max_chunk_size = cls.MAX_CHUNK_SIZE
        if min_chunk_size is None:
            min_chunk_size = cls.MIN_CHUNK_SIZE

        if len(text) <= max_chunk_size:
            return [{
                "content": text,
                "chunk_index": 0,
                "start_char": 0,
                "end_char": len(text),
                "word_count": len(text.split()),
                "is_last_chunk": True
            }]

        chunks = []
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        current_chunk = ""
        current_start = 0
        chunk_index = 0

        for i, paragraph in enumerate(paragraphs):
            # Check if adding this paragraph would exceed max size
            test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph

            if len(test_chunk) > max_chunk_size and len(current_chunk) >= min_chunk_size:
                # Save current chunk and start new one
                chunks.append({
                    "content": current_chunk,
                    "chunk_index": chunk_index,
                    "start_char": current_start,
                    "end_char": current_start + len(current_chunk),
                    "word_count": len(current_chunk.split()),
                    "is_last_chunk": False
                })

                current_chunk = paragraph
                current_start += len(current_chunk) + 2  # +2 for \n\n
                chunk_index += 1
            else:
                # Add paragraph to current chunk
                current_chunk = test_chunk
                if current_chunk == paragraph:
                    current_start = len("".join([c["content"] for c in chunks])) + sum(len(c["content"]) + 2 for c in chunks)

        # Add final chunk
        if current_chunk:
            chunks.append({
                "content": current_chunk,
                "chunk_index": chunk_index,
                "start_char": current_start,
                "end_char": current_start + len(current_chunk),
                "word_count": len(current_chunk.split()),
                "is_last_chunk": True
            })

        logger.info(f"Split text into {len(chunks)} chunks (max size: {max_chunk_size})")
        return chunks

    @classmethod
    def get_optimal_chunk_size(cls, total_length: int, target_chunks: int = 5) -> int:
        """
        Calculate optimal chunk size for given text length and target number of chunks
        """
        if total_length <= cls.MAX_CHUNK_SIZE:
            return total_length

        # Aim for chunks that are not too small or too large
        optimal_size = total_length // target_chunks

        # Ensure chunks are within reasonable bounds
        if optimal_size < cls.MIN_CHUNK_SIZE:
            optimal_size = cls.MIN_CHUNK_SIZE
        elif optimal_size > cls.MAX_CHUNK_SIZE:
            optimal_size = cls.MAX_CHUNK_SIZE

        return optimal_size

    @classmethod
    def cleanup_file(cls, file_path: Path) -> None:
        try:
            if file_path.exists():
                file_path.unlink()
        except Exception:
            pass  # Ignore cleanup errors


